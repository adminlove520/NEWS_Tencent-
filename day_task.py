#用于对网页发送请求
import requests
import datetime
import os
#处理数据用的
from pyquery import PyQuery as pq
#操作word的库
from docx import Document

url='https://i.news.qq.com/trpc.qqnews_web.kv_srv.kv_srv_http_proxy/list?sub_srv_id=24hours&srv_id=pc&offset=1&limit=20&strategy=1&ext={%22pool%22:[%22top%22],%22is_filter%22:7,%22check_type%22:true}'
headers={
  'User-Agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/103.0.0.0 Safari/537.36'
  ,'Cookie':'eas_sid=b136p3q453X1I4h2w1Z5s6D3k2; fqm_pvqid=2912c15f-a1c2-4860-a5fc-cbea585f84da; iip=0; LW_sid=A1i6g3O4x391w4X3O302H1o0n1; LW_uid=f106s3T4S331G482A175G6v2a0; o_cookie=1239676778; pac_uid=1_1239676778; pgv_pvid=6531328166; ptcz=d1c49617e8b14221591f2d667eed59fcbe7ffe2faa4d6cfd6ba76c2134218717; RK=f47AQDUaOd; tvfe_boss_uuid=4a9af621bc6cbdbe; _clck=3085157109|1|f3p|0; pgv_info=ssid=s3679942986; ts_last=news.qq.com/; ts_refer=www.baidu.com/link; ts_uid=3167588418; ad_play_index=38'
  ,'Referer':'https://www.baidu.com/link?url=O8Icaz2zeI7Zb18eJRM034wgs-E4TuaXpV0nJeeaIy3&wd=&eqid=aed05fb000018fd30000000362f39cd3'
}

res=requests.get(url,headers=headers).json()
#从data,list键值中取出数据
data_list=res['data']['list']
#获取当前日期
now = datetime.datetime.now()
#格式化日期
date_str = now.strftime("%Y-%m-%d")
#创建日期文件夹
folder_path = './archive/' + date_str + '/'
os.makedirs(folder_path, exist_ok=True)

#清空archive文件夹下的所有文件
folder_path = './archive/'
for file_name in os.listdir(folder_path):
    file_path = os.path.join(folder_path, file_name)
    if os.path.isfile(file_path):
        os.remove(file_path)
#遍历列表
for data in data_list:
  file = Document()
  #通过cms_id键值取出文章的id
  article_id=data['cms_id']
  #通过"titlt"键值取出文章标题
  article_title=data['title']



  #
  article_url='https://new.qq.com/rain/a/{}'.format(article_id)
  res2=requests.get(article_url).text
  handled_data=pq(res2)
  content_articles=handled_data('.one-p').items()
  for content_article in content_articles:
    news_contents=content_article.text()
    print(news_contents)

    para=file.add_paragraph()
    para.add_run(text="  "+news_contents)
    # 修改为按年-月-日-article_title.docx归档
    # 获取当前日期
    now = datetime.datetime.now()
    # 格式化日期
    date_str = now.strftime("%Y-%m-%d")
    # 保存到本地
    file.save(folder_path + article_title + '.docx')
    #保存到本地
    # file.save('./archive/'+article_title+'.docx')
